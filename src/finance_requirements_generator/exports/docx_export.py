from __future__ import annotations

import re
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from finance_requirements_generator.schemas import RequirementsPack, UATTestCase

CONTENT_DXA = 9360
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_TOP_DXA = 1180
MARGIN_RIGHT_DXA = 1440
MARGIN_BOTTOM_DXA = 1180
MARGIN_LEFT_DXA = 1440

NAVY = "16243D"
BLUE = "2E5C8A"
BLUE_LT = "EAF0F7"
GOLD = "C9A24B"
INK = "1F2430"
GRAY = "6B7280"
ZEBRA = "F5F7FA"
LINE = "D9DEE6"
WHITE = "FFFFFF"
GREEN = "1F7A4D"
GREEN_LT = "E7F1EC"
RED = "9A3B2E"
RED_LT = "F6EBE8"
AMBER = "CFA63A"
AMBER_LT = "FBF5E7"
SUBINK = "3A4253"

CONTENTS_SECTIONS = [
    ("executive_summary", "Executive Summary"),
    ("business_problem", "Business Problem"),
    ("process_scope", "Process Scope"),
    ("scope_boundaries", "Scope Boundaries"),
    ("stakeholders_and_roles", "Stakeholders and Roles"),
    ("functional_requirements", "Functional Requirements"),
    ("data_requirements", "Data Requirements"),
    ("controls", "Controls"),
    ("reporting_requirements", "Reporting Requirements"),
    ("audit_trail_requirements", "Audit Trail Requirements"),
    ("user_stories", "User Stories"),
    ("uat_test_cases", "UAT Test Cases"),
    ("acceptance_criteria", "Acceptance Criteria"),
    ("risks_and_dependencies", "Implementation Risks and Dependencies"),
    ("implementation_notes", "Implementation Notes"),
]

TOC_PAGES = {
    "Executive Summary": "3",
    "Business Problem": "3",
    "Process Scope": "3",
    "Scope Boundaries": "3",
    "Stakeholders and Roles": "4",
    "Functional Requirements": "4",
    "Data Requirements": "5",
    "Controls": "5",
    "Reporting Requirements": "5",
    "Audit Trail Requirements": "5",
    "User Stories": "6",
    "UAT Test Cases": "6",
    "Acceptance Criteria": "7",
    "Implementation Risks and Dependencies": "7",
    "Implementation Notes": "7",
}


def pack_to_docx_bytes(pack: RequirementsPack) -> bytes:
    document = _build_document(pack)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_docx(pack: RequirementsPack, output_path: str | Path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    _build_document(pack).save(path)
    return path


def _build_document(pack: RequirementsPack) -> Document:
    document = Document()
    _configure_document(document)
    _add_header_footer(document, pack)
    _add_cover(document, pack)
    _add_purpose(document)
    _add_contents(document)
    _add_pack_sections(document, pack)
    _add_public_safe_note(document, pack.public_safe_sample_data_note)
    return document


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Twips(PAGE_WIDTH_DXA)
    section.page_height = Twips(PAGE_HEIGHT_DXA)
    section.top_margin = Twips(MARGIN_TOP_DXA)
    section.right_margin = Twips(MARGIN_RIGHT_DXA)
    section.bottom_margin = Twips(MARGIN_BOTTOM_DXA)
    section.left_margin = Twips(MARGIN_LEFT_DXA)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)
    section.different_first_page_header_footer = True

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = _rgb(INK)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = _rgb(NAVY)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    styles["List Paragraph"].font.name = "Arial"
    styles["List Paragraph"].font.size = Pt(9.5)
    styles["List Paragraph"].paragraph_format.space_after = Pt(4)


def _add_header_footer(document: Document, pack: RequirementsPack) -> None:
    section = document.sections[0]
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(3)
    _set_paragraph_border(header, "bottom", LINE, "6")
    _add_tab_stop(header, CONTENT_DXA)
    title_run = header.add_run(f"{pack.process_name} Requirements Pack")
    _style_run(title_run, size=8, bold=True, color=NAVY)
    company_run = header.add_run(f"\t{pack.company_name}")
    _style_run(company_run, size=8, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.paragraph_format.space_before = Pt(2)
    _set_paragraph_border(footer, "top", LINE, "6")
    _add_tab_stop(footer, CONTENT_DXA)
    label_run = footer.add_run("Chez Solutions  \u00b7  Public-safe sample document")
    _style_run(label_run, size=7.5, color=GRAY)
    page_run = footer.add_run("\tPage ")
    _style_run(page_run, size=7.5, color=GRAY)
    _add_page_number(footer)


def _add_cover(document: Document, pack: RequirementsPack) -> None:
    cover = document.add_table(rows=1, cols=1)
    _clear_table_borders(cover)
    _set_table_width(cover, CONTENT_DXA)
    _set_grid_widths(cover, [CONTENT_DXA])
    cell = cover.cell(0, 0)
    _shade_cell(cell, NAVY)
    _set_cell_margins(cell, top=560, bottom=560, left=420, right=420)
    cell.text = ""

    label = cell.paragraphs[0]
    label.paragraph_format.space_after = Pt(6)
    label_run = label.add_run("FINANCE \u00b7 ERP IMPLEMENTATION")
    _style_run(label_run, size=8, bold=True, color=GOLD)

    title = cell.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run(pack.process_name)
    _style_run(title_run, size=28, bold=True, color=WHITE)

    subtitle = cell.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run("Requirements Pack")
    _style_run(subtitle_run, size=28, bold=True, color="C7D3E3")

    rule = cell.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    _set_paragraph_border(rule, "bottom", GOLD, "14")

    prepared = cell.add_paragraph()
    prepared_label = prepared.add_run("Prepared for  ")
    _style_run(prepared_label, size=10, color="9DB0CC")
    prepared_run = prepared.add_run(pack.company_name)
    _style_run(prepared_run, size=11, bold=True, color=WHITE)

    _add_spacer(document, after=7)
    meta = document.add_table(rows=1, cols=4)
    _clear_table_borders(meta)
    _set_table_width(meta, CONTENT_DXA)
    _set_grid_widths(meta, [2340, 2340, 2340, 2340])
    meta_values = [
        ("Prepared by", "Chez Solutions"),
        ("Prepared for", pack.company_name),
        ("Date", date.today().strftime("%B %Y")),
        ("Status", "Draft for Review"),
    ]
    for cell, (label, value) in zip(meta.rows[0].cells, meta_values, strict=True):
        _set_meta_cell(cell, label, value)

    _add_spacer(document, after=8)
    stats = document.add_table(rows=1, cols=5)
    _set_table_width(stats, CONTENT_DXA)
    _set_grid_widths(stats, [1872, 1872, 1872, 1872, 1872])
    _set_table_borders(stats, color=BLUE, size="8", inside_color=LINE, inside_size="4")
    for cell, (value, label) in zip(stats.rows[0].cells, _kpi_values(pack), strict=True):
        _set_kpi_cell(cell, value, label)


def _add_purpose(document: Document) -> None:
    label = document.add_paragraph()
    label.paragraph_format.space_before = Pt(13)
    label.paragraph_format.space_after = Pt(3)
    run = label.add_run("PURPOSE")
    _style_run(run, size=8, bold=True, color=BLUE)

    purpose = document.add_paragraph(
        "Translate finance process pain points into implementation-ready ERP requirements "
        "\u2014 covering workflow, data, controls, reporting, audit trail, and UAT coverage."
    )
    purpose.paragraph_format.space_after = Pt(0)
    purpose.paragraph_format.line_spacing = 1.15
    _set_paragraph_font(purpose, size=11, italic=True, color=_rgb(SUBINK))


def _add_contents(document: Document) -> None:
    heading = document.add_paragraph("Contents")
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.space_after = Pt(3)
    _set_paragraph_font(heading, size=18, bold=True, color=_rgb(NAVY))

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(13)
    _set_paragraph_border(rule, "bottom", GOLD, "12")

    table = document.add_table(rows=len(CONTENTS_SECTIONS), cols=3)
    _clear_table_borders(table)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, [720, 7880, 760])
    for index, (_field_name, title) in enumerate(CONTENTS_SECTIONS, start=1):
        row = table.rows[index - 1]
        _set_toc_cell(row.cells[0], f"{index:02d}", width=720, color=GOLD, bold=True)
        _set_toc_cell(row.cells[1], title, width=7880, color=NAVY, bold=True)
        _set_toc_cell(
            row.cells[2],
            f"p. {TOC_PAGES[title]}",
            width=760,
            color=BLUE,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )


def _add_pack_sections(document: Document, pack: RequirementsPack) -> None:
    _add_numbered_heading(document, 1, "Executive Summary", page_break_before=True)
    document.add_paragraph(pack.executive_summary)

    _add_numbered_heading(document, 2, "Business Problem")
    document.add_paragraph(pack.business_problem)

    _add_numbered_heading(document, 3, "Process Scope")
    document.add_paragraph(pack.process_scope)

    _add_numbered_heading(document, 4, "Scope Boundaries")
    document.add_paragraph(
        "The boundaries below set expectations for this phase. Items out of scope can "
        "be re-planned as separate, approved phases."
    )
    _add_scope_boundaries_table(document, pack)

    _add_numbered_heading(document, 5, "Stakeholders and Roles")
    _add_two_column_table(document, ["Role", "Responsibility"], pack.stakeholders_and_roles)

    _add_numbered_heading(document, 6, "Functional Requirements")
    _add_two_column_table(document, ["ID", "Requirement"], pack.functional_requirements)

    _add_numbered_heading(document, 7, "Data Requirements")
    document.add_paragraph(
        "Minimum data fields required for the controls and reporting outputs to operate reliably."
    )
    _add_data_grid(document, pack.data_requirements)

    _add_numbered_heading(document, 8, "Controls")
    _add_two_column_table(document, ["ID", "Control"], pack.controls)

    _add_numbered_heading(document, 9, "Reporting Requirements")
    _add_two_column_table(document, ["ID", "Requirement"], pack.reporting_requirements)

    _add_numbered_heading(document, 10, "Audit Trail Requirements")
    _add_two_column_table(document, ["ID", "Requirement"], pack.audit_trail_requirements)

    _add_numbered_heading(document, 11, "User Stories")
    _add_user_story_table(document, pack.user_stories)

    _add_numbered_heading(document, 12, "UAT Test Cases")
    _add_uat_table(document, pack.uat_test_cases)

    _add_numbered_heading(document, 13, "Acceptance Criteria")
    document.add_paragraph(
        "The implementation is accepted when the following are demonstrable in test evidence:"
    )
    _add_bullet_list(document, pack.acceptance_criteria, numbered=True)

    _add_numbered_heading(document, 14, "Implementation Risks and Dependencies")
    _add_bullet_list(document, pack.risks_and_dependencies)

    _add_numbered_heading(document, 15, "Implementation Notes")
    _add_bullet_list(document, pack.implementation_notes)


def _add_public_safe_note(document: Document, note: str) -> None:
    _add_spacer(document, after=11)
    table = document.add_table(rows=1, cols=1)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, [CONTENT_DXA])
    _set_table_borders(table, color=AMBER, size="6", left_size="18")
    cell = table.cell(0, 0)
    _shade_cell(cell, AMBER_LT)
    _set_cell_margins(cell, top=150, bottom=150, left=200, right=180)
    cell.text = ""

    label = cell.paragraphs[0]
    label.paragraph_format.space_after = Pt(3)
    label_run = label.add_run("PUBLIC-SAFE SAMPLE DATA NOTE")
    _style_run(label_run, size=8.5, bold=True, color="8A6D1F")

    paragraph = cell.add_paragraph(note)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_paragraph_font(paragraph, size=9.5, color=_rgb(SUBINK))


def _add_scope_boundaries_table(document: Document, pack: RequirementsPack) -> None:
    table = document.add_table(rows=1, cols=2)
    _clear_table_borders(table)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, [4680, 4680])
    _set_scope_cell(table.rows[0].cells[0], "In Scope", pack.in_scope, GREEN, GREEN_LT)
    _set_scope_cell(table.rows[0].cells[1], "Out of Scope", pack.out_of_scope, RED, RED_LT)


def _add_two_column_table(document: Document, headers: list[str], items: list[str]) -> None:
    widths = [1080, 8280] if headers[0] == "ID" else [3000, 6360]
    if headers == ["Role", "User Story"]:
        widths = [2500, 6860]

    table = document.add_table(rows=1, cols=2)
    _style_table(table)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, widths)
    _set_header_row(table.rows[0], headers, widths)
    for row_number, item in enumerate(items, start=1):
        left, right = _split_identifier(item)
        row = table.add_row()
        _set_body_cell(row.cells[0], left, width=widths[0], id_cell=headers[0] == "ID")
        _set_body_cell(row.cells[1], right, width=widths[1], zebra=row_number % 2 == 0)


def _add_data_grid(document: Document, items: list[str]) -> None:
    table = document.add_table(rows=0, cols=4)
    _style_table(table)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, [960, 3720, 960, 3720])
    pairs = [_split_identifier(item) for item in items]
    for row_number, index in enumerate(range(0, len(pairs), 2), start=1):
        row = table.add_row()
        first = pairs[index]
        second = pairs[index + 1] if index + 1 < len(pairs) else ("", "")
        values = [first[0], first[1], second[0], second[1]]
        for col_index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            _set_body_cell(
                cell,
                value,
                width=[960, 3720, 960, 3720][col_index],
                id_cell=col_index in (0, 2),
                zebra=row_number % 2 == 0,
            )


def _add_user_story_table(document: Document, stories: list[str]) -> None:
    table = document.add_table(rows=1, cols=2)
    _style_table(table)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, [2500, 6860])
    _set_header_row(table.rows[0], ["Role", "User Story"], [2500, 6860])
    for row_number, story in enumerate(stories, start=1):
        role, need = _split_user_story(story)
        row = table.add_row()
        _set_body_cell(row.cells[0], role, width=2500, bold=True)
        _set_body_cell(row.cells[1], need, width=6860, zebra=row_number % 2 == 0)


def _add_uat_table(document: Document, cases: list[UATTestCase]) -> None:
    table = document.add_table(rows=1, cols=3)
    _style_table(table)
    _set_table_width(table, CONTENT_DXA)
    _set_grid_widths(table, [880, 4240, 4240])
    _set_header_row(table.rows[0], ["ID", "Test Scenario", "Expected Result"], [880, 4240, 4240])
    for row_number, case in enumerate(cases, start=1):
        row = table.add_row()
        _set_body_cell(row.cells[0], case.test_id, width=880, id_cell=True)
        _set_body_cell(row.cells[1], case.scenario, width=4240, zebra=row_number % 2 == 0)
        _set_body_cell(row.cells[2], case.expected_result, width=4240, zebra=row_number % 2 == 0)


def _add_bullet_list(document: Document, items: list[str], numbered: bool = False) -> None:
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph(style="List Paragraph")
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        prefix = f"{index}. " if numbered else "\u2022 "
        run = paragraph.add_run(f"{prefix}{item}")
        _style_run(run, size=9.5, color=INK, bold=numbered)


def _add_numbered_heading(
    document: Document,
    index: int,
    title: str,
    page_break_before: bool = False,
) -> None:
    heading = document.add_heading(level=1)
    if page_break_before:
        heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(8)
    _set_paragraph_border(heading, "bottom", BLUE, "10")
    number_run = heading.add_run(f"{index:02d}  ")
    _style_run(number_run, size=14, bold=True, color=GOLD)
    title_run = heading.add_run(title)
    _style_run(title_run, size=14, bold=True, color=NAVY)


def _set_meta_cell(cell, label: str, value: str) -> None:
    _shade_cell(cell, ZEBRA)
    _set_cell_margins(cell, top=120, bottom=120, left=150, right=120)
    _set_cell_border(cell, left_color=GOLD, left_size="14", right_color=WHITE, right_size="14")
    cell.text = ""
    label_para = cell.paragraphs[0]
    label_para.paragraph_format.space_after = Pt(1.5)
    label_run = label_para.add_run(label.upper())
    _style_run(label_run, size=6.5, color=GRAY)
    value_para = cell.add_paragraph(value)
    value_para.paragraph_format.space_after = Pt(0)
    _set_paragraph_font(value_para, size=9.5, bold=True, color=_rgb(NAVY))


def _set_kpi_cell(cell, value: str, label: str) -> None:
    _shade_cell(cell, WHITE)
    _set_cell_margins(cell, top=140, bottom=140, left=80, right=80)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    value_para = cell.paragraphs[0]
    value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    value_para.paragraph_format.space_after = Pt(1.5)
    value_run = value_para.add_run(value)
    _style_run(value_run, size=17, bold=True, color=BLUE)
    label_para = cell.add_paragraph(label.upper())
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.space_after = Pt(0)
    _set_paragraph_font(label_para, size=7, color=_rgb(GRAY))


def _set_toc_cell(
    cell,
    text: str,
    width: int,
    color: str,
    bold: bool = False,
    align: int | None = None,
) -> None:
    _set_cell_width(cell, width)
    _set_cell_margins(cell, top=110, bottom=110, left=60, right=60)
    _set_cell_border(cell, bottom_color=LINE, bottom_size="4")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _style_run(run, size=10 if text.startswith("p. ") else 10.5, bold=bold, color=color)


def _set_header_row(row, headers: list[str], widths: list[int]) -> None:
    for cell, header, width in zip(row.cells, headers, widths, strict=True):
        _set_cell_width(cell, width)
        _shade_cell(cell, NAVY)
        _set_cell_margins(cell, top=70, bottom=70, left=130, right=130)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header.upper())
        _style_run(run, size=8, bold=True, color=WHITE)


def _set_body_cell(
    cell,
    text: str,
    width: int,
    id_cell: bool = False,
    zebra: bool = False,
    bold: bool = False,
) -> None:
    _set_cell_width(cell, width)
    _shade_cell(cell, BLUE_LT if id_cell else ZEBRA if zebra else WHITE)
    _set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    _style_run(
        run,
        size=9.5 if id_cell else 10,
        bold=id_cell or bold,
        color=BLUE if id_cell else INK,
    )


def _set_scope_cell(cell, label: str, items: list[str], accent: str, fill: str) -> None:
    _shade_cell(cell, fill)
    _set_cell_margins(cell, top=120, bottom=130, left=160, right=160)
    _set_cell_border(
        cell,
        top_color=accent,
        top_size="8",
        bottom_color=LINE,
        left_color=LINE,
        right_color=LINE,
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    label_para = cell.paragraphs[0]
    _shade_paragraph(label_para, accent)
    label_para.paragraph_format.space_after = Pt(5)
    label_run = label_para.add_run(label.upper())
    _style_run(label_run, size=9, bold=True, color=WHITE)
    for item in items:
        paragraph = cell.add_paragraph(f"\u2022 {item}")
        paragraph.paragraph_format.space_after = Pt(3.5)
        _set_paragraph_font(paragraph, size=9.5, color=_rgb(INK))


def _style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    _set_table_borders(table, color=LINE, size="4")


def _set_table_width(table, width_dxa: int) -> None:
    table.allow_autofit = False
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(width_dxa))
    table_width.set(qn("w:type"), "dxa")


def _set_grid_widths(table, widths: list[int]) -> None:
    for child in list(table._tbl.tblGrid):
        table._tbl.tblGrid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        table._tbl.tblGrid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            _set_cell_width(cell, width)


def _set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.first_child_found_in("w:tcW")
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(width))
    tc_width.set(qn("w:type"), "dxa")


def _set_table_borders(
    table,
    color: str,
    size: str,
    left_size: str | None = None,
    inside_color: str | None = None,
    inside_size: str | None = None,
) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        border_color = inside_color if edge.startswith("inside") and inside_color else color
        border_size = inside_size if edge.startswith("inside") and inside_size else size
        if edge == "left" and left_size:
            border_size = left_size
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), border_size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), border_color)


def _clear_table_borders(table) -> None:
    _set_table_borders(table, color=WHITE, size="0")


def _set_cell_border(
    cell,
    top_color: str | None = None,
    top_size: str = "4",
    bottom_color: str | None = None,
    bottom_size: str = "4",
    left_color: str | None = None,
    left_size: str = "4",
    right_color: str | None = None,
    right_size: str = "4",
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    specs = {
        "top": (top_color, top_size),
        "bottom": (bottom_color, bottom_size),
        "left": (left_color, left_size),
        "right": (right_color, right_size),
    }
    for edge, (color, size) in specs.items():
        if color is None:
            continue
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _shade_paragraph(paragraph, fill: str) -> None:
    paragraph_pr = paragraph._p.get_or_add_pPr()
    shading = paragraph_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        paragraph_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int, bottom: int, left: int, right: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_border(paragraph, edge: str, color: str, size: str) -> None:
    paragraph_pr = paragraph._p.get_or_add_pPr()
    borders = paragraph_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_pr.append(borders)
    element = borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), "4")
    element.set(qn("w:color"), color)


def _add_tab_stop(paragraph, position_dxa: int) -> None:
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Twips(position_dxa), WD_TAB_ALIGNMENT.RIGHT)


def _add_spacer(document: Document, after: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)


def _set_paragraph_font(
    paragraph,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color


def _style_run(
    run,
    size: float,
    color: str,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = _rgb(color)
    run.bold = bold
    run.italic = italic


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def _split_identifier(item: str) -> tuple[str, str]:
    if ": " in item:
        return item.split(": ", 1)
    return "", item


def _split_user_story(story: str) -> tuple[str, str]:
    if not story.startswith("As a ") or ", I want " not in story:
        return "User", story
    role, rest = story[5:].split(", I want ", 1)
    need = rest.replace(" so that ", " so ")
    return role.title(), need[:1].upper() + need[1:]


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.color.rgb = _rgb(NAVY)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _kpi_values(pack: RequirementsPack) -> list[tuple[str, str]]:
    volume_value, volume_label = _volume_tile(pack)
    return [
        (volume_value, volume_label),
        (_delivery_tile(pack), "Delivery Window"),
        (str(len(pack.functional_requirements)), "Functional Reqs"),
        (str(len(pack.controls)), "Key Controls"),
        (str(len(pack.uat_test_cases)), "UAT Cases"),
    ]


def _volume_tile(pack: RequirementsPack) -> tuple[str, str]:
    volume = _pack_volume(pack)
    process_labels = {
        "accounts_payable": "Invoices / mo",
        "accounts_receivable": "Invoices / mo",
        "bank_reconciliation": "Statement lines",
        "vat_reconciliation": "Tax codes",
        "month_end_close": "Close tasks",
        "inventory_costing": "SKUs",
        "intercompany_settlements": "Entities",
        "payroll_controls": "Employees",
    }
    number = re.search(r"\d[\d,]*", volume)
    value = number.group(0) if number else volume
    return value, process_labels.get(pack.process_key, "Volume")


def _pack_volume(pack: RequirementsPack) -> str:
    for line in pack.executive_summary.split(". "):
        marker = "It is sized for "
        if marker in line:
            return line.split(marker, 1)[1].split(" and frames", 1)[0]
    return "Defined"


def _delivery_tile(pack: RequirementsPack) -> str:
    window = _extract_delivery_window(pack)
    match = re.search(r"(\d+)\s+weeks?", window, re.IGNORECASE)
    if match:
        return f"{match.group(1)} wks"
    return window


def _extract_delivery_window(pack: RequirementsPack) -> str:
    marker = "target delivery window of "
    if marker in pack.executive_summary:
        return pack.executive_summary.split(marker, 1)[1].rstrip(".")
    return "Draft"
