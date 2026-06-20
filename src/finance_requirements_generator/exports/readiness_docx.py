from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from finance_requirements_generator.exports.readiness_markdown import _traceability_rows
from finance_requirements_generator.schemas import (
    ImplementationReadinessPack,
    OpenDecision,
    ReadinessCheck,
    WorkshopQuestion,
)

NAVY = "0F1B2D"
GREEN = "0E7C66"
GOLD = "C9972B"
LINE = "DDE3EA"
PALE = "F4F7FA"
WHITE = "FFFFFF"


def pack_to_readiness_docx_bytes(pack: ImplementationReadinessPack) -> bytes:
    document = _build_document(pack)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_readiness_docx(
    pack: ImplementationReadinessPack,
    output_path: str | Path,
) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(pack_to_readiness_docx_bytes(pack))
    return path


def _build_document(pack: ImplementationReadinessPack) -> Document:
    document = Document()
    _configure_document(document)
    _add_header_footer(document, pack)
    _add_cover(document, pack)
    document.add_section(WD_SECTION.NEW_PAGE)
    _add_heading(document, "1. Readiness Summary", level=1)
    document.add_paragraph(pack.readiness_summary)
    _add_check_section(
        document,
        "2. Process Implementation Checklist",
        pack.process_implementation_checklist,
    )
    _add_check_section(
        document,
        "3. Target-System Readiness",
        pack.target_system_readiness_checklist,
    )
    _add_check_section(document, "4. Data Readiness", pack.data_readiness_checklist)
    _add_check_section(
        document,
        "5. Controls and UAT Readiness",
        pack.controls_and_uat_readiness_checklist,
    )
    _add_workshop_section(document, pack.configuration_workshop_questions)
    _add_check_section(document, "7. Cutover Readiness", pack.cutover_readiness_notes)
    _add_decisions_section(document, pack.open_decisions_and_dependencies)
    _add_traceability(document, pack)
    _add_heading(document, "10. Public-Safe Sample Data Note", level=1)
    note = document.add_paragraph(pack.public_safe_sample_data_note)
    note.style = document.styles["Intense Quote"]
    return document


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    for style_name, size in (("Title", 30), ("Heading 1", 16), ("Heading 2", 12)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = _rgb(NAVY)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True


def _add_header_footer(document: Document, pack: ImplementationReadinessPack) -> None:
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = f"{pack.process_name} Implementation Readiness Pack | {pack.company_name}"
        _style_paragraph(header, 8, NAVY)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Chez Solutions | Public-safe sample document | Page ")
        _add_page_number(footer)
        _style_paragraph(footer, 8, "6B7280")


def _add_cover(document: Document, pack: ImplementationReadinessPack) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade_cell(cell, NAVY)
    cell.text = ""
    label = cell.paragraphs[0]
    label.add_run("FINANCE · ERP IMPLEMENTATION READINESS").bold = True
    _style_paragraph(label, 9, GOLD)
    title = cell.add_paragraph(pack.process_name)
    _style_paragraph(title, 28, WHITE, bold=True)
    subtitle = cell.add_paragraph("Implementation Readiness Pack")
    _style_paragraph(subtitle, 24, "C7D3E3", bold=True)
    company = cell.add_paragraph(f"Prepared for {pack.company_name}")
    _style_paragraph(company, 12, WHITE, bold=True)
    target = cell.add_paragraph(f"Target system: {pack.target_system}")
    _style_paragraph(target, 10, "C7D3E3")
    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(18)
    meta.add_run(
        f"Prepared by Chez Solutions    |    {date.today().strftime('%B %Y')}    |    "
        "Status: Draft for Review"
    )
    _style_paragraph(meta, 10, NAVY, bold=True)
    purpose = document.add_paragraph(
        "Purpose: determine whether the finance process is ready to configure, test, migrate, "
        "and implement using traceable requirements, controls, data, UAT, risks, and "
        "fit-gap inputs."
    )
    purpose.paragraph_format.space_before = Pt(14)
    _style_paragraph(purpose, 11, "3A4659")


def _add_check_section(
    document: Document,
    title: str,
    checks: list[ReadinessCheck],
) -> None:
    _add_heading(document, title, level=1)
    table = document.add_table(rows=1, cols=6)
    _set_headers(table, ["ID", "Finance-Specific Check", "Evidence", "Owner", "Status", "Sources"])
    for item in checks:
        row = table.add_row().cells
        values = [
            item.check_id,
            f"{item.finance_specific_check}\n\nValidation: {item.validation_note}",
            item.evidence_required,
            item.suggested_owner_role,
            item.review_status,
            ", ".join(item.source_references),
        ]
        _set_row(row, values, [0.6, 2.25, 1.7, 1.15, 0.75, 1.05])


def _add_workshop_section(document: Document, items: list[WorkshopQuestion]) -> None:
    _add_heading(document, "6. Configuration Workshop Questions", level=1)
    table = document.add_table(rows=1, cols=4)
    _set_headers(table, ["ID", "Question", "Implementation Relevance", "Sources"])
    for item in items:
        _set_row(
            table.add_row().cells,
            [
                item.question_id,
                item.question,
                item.implementation_relevance,
                ", ".join(item.source_references),
            ],
            [0.7, 2.8, 2.75, 1.2],
        )


def _add_decisions_section(document: Document, items: list[OpenDecision]) -> None:
    _add_heading(document, "8. Open Decisions and Dependencies", level=1)
    table = document.add_table(rows=1, cols=5)
    _set_headers(table, ["ID", "Decision Required", "Dependency / Impact", "Owner", "Sources"])
    for item in items:
        _set_row(
            table.add_row().cells,
            [
                item.decision_id,
                item.decision_required,
                item.dependency_or_impact,
                item.suggested_decision_owner,
                ", ".join(item.source_references),
            ],
            [0.65, 2.25, 2.2, 1.3, 1.05],
        )


def _add_traceability(document: Document, pack: ImplementationReadinessPack) -> None:
    _add_heading(document, "9. Source Traceability", level=1)
    document.add_paragraph(
        "References connect readiness prompts to generated requirements, controls, reports, "
        "audit needs, UAT cases, risks, fit-gap rows, and control-risk matrix rows."
    )
    table = document.add_table(rows=1, cols=2)
    _set_headers(table, ["Readiness Item", "Source References"])
    for item_id, references in _traceability_rows(pack):
        _set_row(table.add_row().cells, [item_id, ", ".join(references)], [1.3, 6.2])


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(5)


def _set_headers(table, headers: list[str]) -> None:
    table.style = "Table Grid"
    table.autofit = False
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
        _shade_cell(cell, NAVY)
        _style_paragraph(cell.paragraphs[0], 8, WHITE, bold=True)
        _repeat_table_header(table.rows[0])


def _set_row(cells, values: list[str], widths: list[float]) -> None:
    for index, (cell, value, width) in enumerate(zip(cells, values, widths, strict=True)):
        cell.text = value
        cell.width = Inches(width)
        _style_paragraph(cell.paragraphs[0], 8, NAVY, bold=index == 0)
        if index == 0:
            _prevent_wrap(cell)
        if index % 2 == 0:
            _shade_cell(cell, PALE)


def _style_paragraph(
    paragraph,
    size: float,
    color: str,
    bold: bool = False,
) -> None:
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = _rgb(color)


def _shade_cell(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _prevent_wrap(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    if properties.find(qn("w:noWrap")) is None:
        properties.append(OxmlElement("w:noWrap"))


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        properties.append(OxmlElement("w:tblHeader"))


def _add_page_number(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, end])


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)
