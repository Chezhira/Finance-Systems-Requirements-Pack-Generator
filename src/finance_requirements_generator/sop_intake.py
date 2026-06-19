from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from docx import Document

from finance_requirements_generator.schemas import IntakeAnswers, SOPDraft

SECTION_ALIASES = {
    "trigger": ("trigger", "start", "when"),
    "owner": ("owner", "process owner", "responsible"),
    "systems": ("systems", "tools", "applications"),
    "steps": ("steps", "procedure", "process steps"),
    "approvals": ("approval", "approvals", "authorisation", "authorization"),
    "handoffs": ("handoff", "handoffs", "handover"),
    "controls": ("control", "controls"),
    "reports": ("report", "reports", "reporting"),
    "exceptions": ("exception", "exceptions", "escalation"),
    "data_fields": ("data", "fields", "data fields"),
    "audit_evidence": ("audit", "evidence"),
    "pain_points": ("pain", "risk", "issue", "gap"),
    "future_improvements": ("future", "improvement", "target state"),
}


@dataclass(frozen=True)
class GuidedSOPAnswers:
    process_trigger: str
    process_owner: str
    systems_tools: str
    key_process_steps: list[str]
    approvals: list[str]
    handoffs: list[str]
    controls: list[str]
    reports: list[str]
    exceptions: list[str]
    data_fields: list[str]
    audit_evidence: list[str]
    pain_points: list[str]
    desired_future_state_improvements: list[str]


@dataclass(frozen=True)
class SOPMappedFields:
    current_tools: str
    pain_points: list[str]
    control_concerns: list[str]
    reporting_needs: str
    compliance_focus: str
    assumptions: list[str] = field(default_factory=list)
    sop_draft: SOPDraft | None = None


def extract_text_from_upload(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return content.decode("utf-8", errors="ignore")
    if suffix == ".docx":
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(item for item in paragraphs + table_cells if item.strip())
    raise ValueError("Only TXT, Markdown, and DOCX SOP uploads are supported in v0.3.0.")


def map_sop_text_to_intake(text: str, fallback: IntakeAnswers) -> SOPMappedFields:
    sections = _section_buckets(text)
    current_tools = _first_nonempty(
        _join_lines(sections["systems"]),
        fallback.current_tools,
    )
    controls = _items_from_lines(sections["controls"]) or fallback.control_concerns
    reports = _items_from_lines(sections["reports"])
    pain_points = _items_from_lines(sections["pain_points"]) or fallback.pain_points
    evidence = _items_from_lines(sections["audit_evidence"])
    approvals = _items_from_lines(sections["approvals"])

    sop_draft = SOPDraft(
        purpose=(
            "Document the current-state finance process so requirements can be reviewed "
            "before ERP or finance-system implementation decisions are made."
        ),
        scope=fallback.entity_type,
        trigger=_first_nonempty(_join_lines(sections["trigger"]), "Trigger to be confirmed."),
        roles_and_responsibilities=_first_nonempty(
            _join_lines(sections["owner"]),
            f"{fallback.sponsor} to confirm process ownership and sign-off.",
        ),
        step_by_step_procedure=_items_from_lines(sections["steps"])
        or ["Process steps to confirm."],
        controls_and_approvals=controls + approvals,
        exceptions_and_escalations=_items_from_lines(sections["exceptions"]) or [
            "Exceptions and escalations to confirm."
        ],
        reports_and_evidence=reports + evidence,
        systems_and_data_used=_items_from_lines(sections["systems"] + sections["data_fields"])
        or [fallback.current_tools],
        review_and_sign_off=(
            "Finance process owner should review extracted fields, confirm gaps, and sign off "
            "before pack generation."
        ),
    )

    return SOPMappedFields(
        current_tools=current_tools,
        pain_points=pain_points,
        control_concerns=controls,
        reporting_needs=_first_nonempty(_join_lines(sections["reports"]), fallback.reporting_needs),
        compliance_focus=_first_nonempty(
            _join_lines(sections["controls"] + sections["audit_evidence"]),
            fallback.compliance_focus,
        ),
        assumptions=["Current-state SOP fields were mapped deterministically and require review."],
        sop_draft=sop_draft,
    )


def build_sop_draft(answers: GuidedSOPAnswers, process_name: str, entity_type: str) -> SOPDraft:
    return SOPDraft(
        purpose=(
            f"Document the current-state {process_name} process for review before requirements "
            "pack generation."
        ),
        scope=f"{process_name} activities for {entity_type}.",
        trigger=answers.process_trigger,
        roles_and_responsibilities=answers.process_owner,
        step_by_step_procedure=answers.key_process_steps,
        controls_and_approvals=answers.controls + answers.approvals,
        exceptions_and_escalations=answers.exceptions + answers.handoffs,
        reports_and_evidence=answers.reports + answers.audit_evidence,
        systems_and_data_used=[answers.systems_tools, *answers.data_fields],
        review_and_sign_off=(
            "Review the SOP draft, update unclear handoffs or evidence gaps, and confirm owner "
            "sign-off before generating the requirements pack."
        ),
    )


def sop_draft_to_text(sop: SOPDraft) -> str:
    sections: list[tuple[str, str | list[str]]] = [
        ("Purpose", sop.purpose),
        ("Scope", sop.scope),
        ("Trigger", sop.trigger),
        ("Roles and Responsibilities", sop.roles_and_responsibilities),
        ("Step-by-Step Procedure", sop.step_by_step_procedure),
        ("Controls and Approvals", sop.controls_and_approvals),
        ("Exceptions and Escalations", sop.exceptions_and_escalations),
        ("Reports and Evidence", sop.reports_and_evidence),
        ("Systems and Data Used", sop.systems_and_data_used),
        ("Review and Sign-off", sop.review_and_sign_off),
    ]
    lines = []
    for title, value in sections:
        lines.append(f"## {title}")
        if isinstance(value, list):
            lines.extend(f"- {item}" for item in value if item)
        else:
            lines.append(value)
        lines.append("")
    return "\n".join(lines).strip()


def guided_answers_to_mapped_fields(
    answers: GuidedSOPAnswers,
    process_name: str,
    entity_type: str,
) -> SOPMappedFields:
    sop = build_sop_draft(answers, process_name, entity_type)
    return SOPMappedFields(
        current_tools=answers.systems_tools,
        pain_points=answers.pain_points,
        control_concerns=answers.controls,
        reporting_needs=", ".join(answers.reports) or "Reporting needs to confirm.",
        compliance_focus=", ".join(answers.controls + answers.audit_evidence)
        or "Control evidence needs to confirm.",
        assumptions=["Guided SOP draft should be reviewed by the finance process owner."],
        sop_draft=sop,
    )


def _section_buckets(text: str) -> dict[str, list[str]]:
    buckets = {key: [] for key in SECTION_ALIASES}
    current_key = ""
    for raw_line in text.splitlines():
        line = raw_line.strip(" \t#-*:")
        if not line:
            continue
        detected = _detect_section(line)
        if detected:
            current_key = detected
            continue
        if current_key:
            buckets[current_key].append(line)
        else:
            lowered = line.lower()
            for key, aliases in SECTION_ALIASES.items():
                if any(alias in lowered for alias in aliases):
                    buckets[key].append(line)
                    break
    return buckets


def _detect_section(line: str) -> str:
    lowered = line.lower()
    if len(line.split()) > 8:
        return ""
    for key, aliases in SECTION_ALIASES.items():
        if any(alias == lowered or lowered.startswith(f"{alias} ") for alias in aliases):
            return key
    return ""


def _items_from_lines(lines: list[str]) -> list[str]:
    items = []
    for line in lines:
        for piece in line.replace(";", "\n").split("\n"):
            item = piece.strip(" -*")
            if item:
                items.append(item)
    return items[:8]


def _join_lines(lines: list[str]) -> str:
    return "; ".join(_items_from_lines(lines))


def _first_nonempty(*values: str) -> str:
    for value in values:
        if value and value.strip():
            return value.strip()
    return ""
