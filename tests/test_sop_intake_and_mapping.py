from io import BytesIO

from docx import Document

from finance_requirements_generator.exports import pack_to_docx_bytes, pack_to_markdown
from finance_requirements_generator.questionnaire import DEFAULT_SAMPLE_INPUTS
from finance_requirements_generator.schemas import IntakeAnswers
from finance_requirements_generator.sop_intake import (
    GuidedSOPAnswers,
    extract_text_from_upload,
    guided_answers_to_mapped_fields,
    map_sop_text_to_intake,
)
from finance_requirements_generator.system_mapping import get_fit_gap_mapping
from finance_requirements_generator.template_engine import generate_pack

SOP_TEXT = """
# Accounts Payable SOP

Trigger
Supplier invoice received in the shared AP mailbox.

Systems
ERP purchase ledger; AP tracker; approval email.

Controls
Duplicate invoice review
Approval threshold check
Supplier bank detail review

Reports
Aged approval backlog and blocked invoices report.

Pain Points
Manual approval chasing
Duplicate invoice risk

Audit Evidence
Invoice image, approver, timestamp, and payment batch reference.
"""


def test_txt_markdown_and_docx_text_extraction() -> None:
    assert "Supplier invoice" in extract_text_from_upload("ap.md", SOP_TEXT.encode())
    assert "Supplier invoice" in extract_text_from_upload("ap.txt", SOP_TEXT.encode())

    document = Document()
    document.add_paragraph("Payroll Controls SOP")
    document.add_paragraph("Controls")
    document.add_paragraph("Payroll change approval")
    buffer = BytesIO()
    document.save(buffer)

    extracted = extract_text_from_upload("payroll.docx", buffer.getvalue())
    assert "Payroll Controls SOP" in extracted
    assert "Payroll change approval" in extracted


def test_deterministic_sop_mapping_populates_reviewable_fields() -> None:
    sample = DEFAULT_SAMPLE_INPUTS["accounts_payable"]
    mapped = map_sop_text_to_intake(SOP_TEXT, sample)

    assert "ERP purchase ledger" in mapped.current_tools
    assert "Manual approval chasing" in mapped.pain_points
    assert "Duplicate invoice review" in mapped.control_concerns
    assert "Aged approval backlog" in mapped.reporting_needs
    assert mapped.sop_draft is not None
    assert "Supplier invoice received" in mapped.sop_draft.trigger


def test_guided_sop_draft_generation() -> None:
    answers = GuidedSOPAnswers(
        process_trigger="Close period starts on working day one.",
        process_owner="Financial Controller",
        systems_tools="Close checklist, ERP journals, and evidence folder",
        key_process_steps=["Open close checklist", "Post journals", "Review reconciliations"],
        approvals=["Controller sign-off"],
        handoffs=["Finance preparer to controller"],
        controls=["Journal approval", "Reconciliation review"],
        reports=["Close status report"],
        exceptions=["Late reconciliation"],
        data_fields=["Period", "Owner", "Status"],
        audit_evidence=["Journal approval evidence"],
        pain_points=["Late close tasks"],
        desired_future_state_improvements=["Automated overdue task escalation"],
    )

    mapped = guided_answers_to_mapped_fields(answers, "Month-end Close", "Retail group")

    assert mapped.sop_draft is not None
    assert mapped.sop_draft.scope == "Month-end Close activities for Retail group."
    assert "Journal approval" in mapped.control_concerns
    assert "Close status report" in mapped.reporting_needs


def test_target_system_mapping_lookup_is_curated_and_cautious() -> None:
    rows = get_fit_gap_mapping("payroll_controls", "Xero")

    assert rows
    assert rows[0].current_state_area
    assert "Potential" in rows[0].candidate_fit_gap_view
    assert "validate" in rows[0].validation_note.lower()


def test_pack_includes_target_mapping_and_exports() -> None:
    sample = DEFAULT_SAMPLE_INPUTS["payroll_controls"]
    intake = IntakeAnswers(
        **{
            **sample.__dict__,
            "target_system": "Xero",
        }
    )
    pack = generate_pack(intake)
    markdown = pack_to_markdown(pack)
    docx = Document(BytesIO(pack_to_docx_bytes(pack)))
    paragraph_text = [paragraph.text for paragraph in docx.paragraphs]
    table_text = [
        cell.text
        for table in docx.tables
        for row in table.rows
        for cell in row.cells
    ]
    text = "\n".join(paragraph_text + table_text)

    assert pack.target_system_fit_gap_mapping
    assert "## Target-System Fit-Gap Mapping" in markdown
    assert "Candidate mapping only" in markdown
    assert "Target-System Fit-Gap Mapping" in text
    assert "Public-Safe Sample Data Note" in markdown
    assert markdown.rfind("Public-Safe Sample Data Note") > markdown.rfind(
        "Target-System Fit-Gap Mapping"
    )
