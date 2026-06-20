from dataclasses import replace
from html import escape
from io import BytesIO

from docx import Document

from finance_requirements_generator.exports import (
    pack_to_readiness_docx_bytes,
    pack_to_readiness_markdown,
)
from finance_requirements_generator.process_library import load_all_templates
from finance_requirements_generator.questionnaire import DEFAULT_SAMPLE_INPUTS
from finance_requirements_generator.readiness_engine import (
    generate_implementation_readiness_pack,
    readiness_source_references,
)
from finance_requirements_generator.template_engine import generate_pack

REQUIRED_READINESS_HEADINGS = [
    "Readiness Summary",
    "Process Implementation Checklist",
    "Target-System Readiness",
    "Data Readiness",
    "Controls and UAT Readiness",
    "Configuration Workshop Questions",
    "Cutover Readiness",
    "Open Decisions and Dependencies",
    "Source Traceability",
    "Public-Safe Sample Data Note",
]


def test_all_templates_have_valid_process_specific_readiness_configuration() -> None:
    templates = load_all_templates()
    expected_terms = {
        "accounts_payable": "invoice",
        "bank_reconciliation": "bank",
        "vat_reconciliation": "vat",
        "accounts_receivable": "receipt",
        "month_end_close": "close",
        "inventory_costing": "costing",
        "intercompany_settlements": "entity",
        "payroll_controls": "payroll",
    }

    for process_key, template in templates.items():
        readiness = template["implementation_readiness"]
        assert len(readiness["process_checks"]) >= 2
        assert len(readiness["configuration_workshop_questions"]) >= 3
        assert len(readiness["cutover_considerations"]) >= 2
        rendered = str(readiness).lower()
        assert expected_terms[process_key] in rendered


def test_all_seven_readiness_areas_generate_with_unique_ids_and_resolved_sources() -> None:
    for intake in DEFAULT_SAMPLE_INPUTS.values():
        requirements_pack = generate_pack(intake)
        readiness_pack = generate_implementation_readiness_pack(requirements_pack)
        groups = [
            readiness_pack.process_implementation_checklist,
            readiness_pack.target_system_readiness_checklist,
            readiness_pack.data_readiness_checklist,
            readiness_pack.controls_and_uat_readiness_checklist,
            readiness_pack.configuration_workshop_questions,
            readiness_pack.cutover_readiness_notes,
            readiness_pack.open_decisions_and_dependencies,
        ]
        assert all(groups)

        item_ids = _readiness_ids(readiness_pack)
        assert len(item_ids) == len(set(item_ids))
        assert all(item.review_status == "Not assessed" for group in groups[:4] for item in group)
        assert all(
            item.review_status == "Not assessed"
            for item in readiness_pack.cutover_readiness_notes
        )

        valid_sources = _valid_source_ids(requirements_pack)
        assert readiness_source_references(readiness_pack).issubset(valid_sources)


def test_selected_target_system_mapping_is_cautious_and_unselected_path_is_explicit() -> None:
    selected_intake = replace(
        DEFAULT_SAMPLE_INPUTS["accounts_payable"],
        target_system="Odoo",
    )
    selected = generate_pack(selected_intake)
    selected_readiness = generate_implementation_readiness_pack(selected)
    selected_text = str(selected_readiness.target_system_readiness_checklist).lower()

    assert selected.target_system
    assert "candidate mapping only" in selected_text
    assert "implementation validation" in selected_text
    assert "edition" in selected_text
    assert "localisation" in selected_text

    no_target_intake = replace(
        selected_intake,
        target_system="",
    )
    no_target = generate_implementation_readiness_pack(generate_pack(no_target_intake))
    assert no_target.target_system == "Not selected"
    assert "target system" in no_target.target_system_readiness_checklist[0].finance_specific_check
    assert no_target.open_decisions_and_dependencies[0].decision_id == "DEC-00"

    selected_without_mapping = generate_implementation_readiness_pack(
        generate_pack(
            replace(
                DEFAULT_SAMPLE_INPUTS["inventory_costing"],
                target_system="Odoo",
            )
        )
    )
    assert selected_without_mapping.target_system == "Odoo"
    assert "no curated process mapping" in (
        selected_without_mapping.target_system_readiness_checklist[0].validation_note.lower()
    )
    assert all(
        item.decision_id != "DEC-00"
        for item in selected_without_mapping.open_decisions_and_dependencies
    )


def test_readiness_generation_is_deterministic_and_does_not_use_network(monkeypatch) -> None:
    def fail_network(*args, **kwargs):
        raise AssertionError("Readiness generation must not access the network")

    monkeypatch.setattr("socket.create_connection", fail_network)
    requirements_pack = generate_pack(DEFAULT_SAMPLE_INPUTS["month_end_close"])

    first = generate_implementation_readiness_pack(requirements_pack)
    second = generate_implementation_readiness_pack(requirements_pack)
    assert first == second


def test_control_readiness_uses_semantically_aligned_uat_cases() -> None:
    requirements_pack = generate_pack(DEFAULT_SAMPLE_INPUTS["accounts_payable"])
    readiness_pack = generate_implementation_readiness_pack(requirements_pack)
    references = {
        item.check_id: item.source_references
        for item in readiness_pack.controls_and_uat_readiness_checklist
    }

    assert "UAT-01" in references["CU-01"]
    assert "UAT-03" in references["CU-02"]
    assert "UAT-04" in references["CU-03"]
    assert "UAT-02" in references["CU-04"]
    assert all(not reference.startswith("UAT-") for reference in references["CU-05"])


def test_readiness_markdown_and_docx_exports_contain_required_sections() -> None:
    requirements_pack = generate_pack(DEFAULT_SAMPLE_INPUTS["payroll_controls"])
    readiness_pack = generate_implementation_readiness_pack(requirements_pack)
    markdown = pack_to_readiness_markdown(readiness_pack)
    docx_bytes = pack_to_readiness_docx_bytes(readiness_pack)
    document = Document(BytesIO(docx_bytes))
    document_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    assert docx_bytes.startswith(b"PK")
    assert "Not assessed" in markdown
    assert "Not assessed" in document_text
    for heading in REQUIRED_READINESS_HEADINGS:
        assert f"## {heading}" in markdown
        assert heading in document_text
    assert escape("Payroll Controls") in markdown
    assert len(document.tables) >= 9


def _readiness_ids(pack) -> list[str]:
    ids = []
    for group in (
        pack.process_implementation_checklist,
        pack.target_system_readiness_checklist,
        pack.data_readiness_checklist,
        pack.controls_and_uat_readiness_checklist,
        pack.cutover_readiness_notes,
    ):
        ids.extend(item.check_id for item in group)
    ids.extend(item.question_id for item in pack.configuration_workshop_questions)
    ids.extend(item.decision_id for item in pack.open_decisions_and_dependencies)
    return ids


def _valid_source_ids(pack) -> set[str]:
    identifiers = {"INTAKE-TARGET-SYSTEM"}
    for group in (
        pack.functional_requirements,
        pack.data_requirements,
        pack.controls,
        pack.reporting_requirements,
        pack.audit_trail_requirements,
    ):
        identifiers.update(item.split(":", 1)[0] for item in group)
    identifiers.update(case.test_id for case in pack.uat_test_cases)
    identifiers.update(
        f"RISK-{index:02d}" for index, _item in enumerate(pack.risks_and_dependencies, start=1)
    )
    identifiers.update(
        f"FG-{index:02d}"
        for index, _item in enumerate(pack.target_system_fit_gap_mapping, start=1)
    )
    identifiers.update(
        f"CRM-{index:02d}" for index, _item in enumerate(pack.control_risk_matrix, start=1)
    )
    return identifiers
