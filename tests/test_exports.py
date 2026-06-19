from io import BytesIO
from zipfile import ZipFile

from docx import Document

from finance_requirements_generator.exports import (
    export_docx,
    export_markdown,
    pack_to_docx_bytes,
    pack_to_markdown,
)
from finance_requirements_generator.questionnaire import DEFAULT_SAMPLE_INPUTS
from finance_requirements_generator.template_engine import generate_pack


def test_markdown_export_contains_key_headings(tmp_path) -> None:
    for process_key, intake in DEFAULT_SAMPLE_INPUTS.items():
        pack = generate_pack(intake)
        output_path = export_markdown(pack, tmp_path / f"{process_key}_pack.md")

        content = output_path.read_text(encoding="utf-8")
        assert f"# {pack.process_name} Requirements Pack" in content
        assert "## Functional Requirements" in content
        assert "## UAT Test Cases" in content
        assert pack_to_markdown(pack) == content


def test_markdown_export_leads_with_business_problem_not_sample_note() -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["accounts_payable"])
    content = pack_to_markdown(pack)
    top_block = "\n".join(content.splitlines()[:12]).lower()

    assert "## business problem" in content.lower()
    assert "## reporting requirements" in content.lower()
    assert "## public-safe sample data note" in content.lower()
    assert "synthetic" not in top_block
    assert "fictional" not in top_block
    assert content.index("## Reporting Requirements") < content.index("## Audit Trail Requirements")
    assert content.index("## Public-Safe Sample Data Note") > content.index(
        "## Implementation Notes"
    )


def test_docx_export_contains_key_headings(tmp_path) -> None:
    for process_key, intake in DEFAULT_SAMPLE_INPUTS.items():
        pack = generate_pack(intake)
        output_path = export_docx(pack, tmp_path / f"{process_key}_pack.docx")

        document = Document(output_path)
        text = _docx_text(document)
        assert pack.process_name in text
        assert "Requirements Pack" in text
        assert "Functional Requirements" in text
        assert "Reporting Requirements" in text
        assert "PUBLIC-SAFE SAMPLE DATA NOTE" in text
        assert "UAT Test Cases" in text


def test_docx_export_matches_reference_cover_contents_and_pagination(tmp_path) -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["accounts_payable"])
    output_path = export_docx(pack, tmp_path / "ap_pack.docx")
    document = Document(output_path)
    text = _docx_text(document)
    section = document.sections[0]

    with ZipFile(output_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        header_footer_parts = [
            item
            for item in archive.namelist()
            if item.startswith("word/header") or item.startswith("word/footer")
        ]

    assert section.page_width.twips == 12240
    assert section.page_height.twips == 15840
    assert section.top_margin.twips == 1180
    assert section.right_margin.twips == 1440
    assert section.bottom_margin.twips == 1180
    assert section.left_margin.twips == 1440
    assert section.different_first_page_header_footer

    assert "FINANCE · ERP IMPLEMENTATION" in text
    assert "Prepared for" in document.tables[0].cell(0, 0).text
    assert "PREPARED BY" in document.tables[1].cell(0, 0).text
    assert "Chez Solutions" in document.tables[1].cell(0, 0).text
    assert "1,200\nINVOICES / MO" in document.tables[2].cell(0, 0).text

    assert "Contents" in text
    assert "01\nExecutive Summary\np. 3" in _table_text(document.tables[3])
    assert "15\nImplementation Notes\np. 7" in _table_text(document.tables[3])
    assert len(document.tables) >= 10

    assert xml.count('w:type="page"') == 0
    assert "w:pageBreakBefore" in xml
    assert "{[object Object]-0}" not in xml
    assert xml.count("<w:noWrap") >= 30
    assert '<w:tcW w:type="dxa" w:w="1250"/>' in xml
    assert header_footer_parts
    assert document.tables[-1].cell(0, 0).text.endswith(
        "Do not upload confidential business information into a public demo."
    )


def test_docx_download_bytes_open_as_word_document() -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["payroll_controls"])
    content = pack_to_docx_bytes(pack)
    document = Document(BytesIO(content))
    text = _docx_text(document)

    assert content.startswith(b"PK")
    assert "Payroll Controls" in text


def _docx_text(document: Document) -> str:
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    table_text = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraph_text + table_text)


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)
