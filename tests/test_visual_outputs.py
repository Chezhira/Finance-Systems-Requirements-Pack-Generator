from html import escape
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from finance_requirements_generator.exports import (
    pack_control_risk_csv_bytes,
    pack_control_risk_xlsx_bytes,
    pack_process_map_html_bytes,
    pack_to_docx_bytes,
    pack_to_markdown,
)
from finance_requirements_generator.questionnaire import DEFAULT_SAMPLE_INPUTS
from finance_requirements_generator.template_engine import generate_pack
from scripts.generate_examples import generate_examples


def test_control_risk_matrix_generation() -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["accounts_payable"])
    row = pack.control_risk_matrix[0]

    assert pack.control_risk_matrix
    assert row.process_area == "Accounts Payable"
    assert row.control_type in {
        "Preventive",
        "Detective",
        "Corrective",
        "Manual",
        "Automated",
        "Semi-automated",
    }
    assert row.related_requirement_id.startswith("FR-")
    assert row.related_uat_case.startswith("UAT-")


def test_control_risk_csv_and_xlsx_exports_open() -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["payroll_controls"])
    csv_content = pack_control_risk_csv_bytes(pack).decode("utf-8-sig")
    xlsx_content = pack_control_risk_xlsx_bytes(pack)

    assert "Process Area,Risk Area,Risk Description" in csv_content
    workbook = load_workbook(BytesIO(xlsx_content))
    worksheet = workbook["Control Risk Matrix"]
    assert worksheet.freeze_panes == "A2"
    assert worksheet["A1"].value == "Process Area"
    assert worksheet["A2"].value == "Payroll Controls"
    assert worksheet.column_dimensions["A"].width >= 20


def test_mermaid_process_map_generation() -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["month_end_close"])

    assert pack.mermaid_process_map.startswith("flowchart LR")
    assert "Exception or control gap?" in pack.mermaid_process_map
    assert "G[" in pack.mermaid_process_map
    assert "dashboard" in pack.mermaid_process_map
    assert "class A,H gate" in pack.mermaid_process_map
    assert "class B,C,F,G step" in pack.mermaid_process_map
    assert "class D decide" in pack.mermaid_process_map
    assert "class E fix" in pack.mermaid_process_map
    assert pack.process_map_summary


def test_browser_process_map_export_is_styled_and_offline_safe(monkeypatch) -> None:
    def fail_network(*args, **kwargs):
        raise AssertionError("HTML export generation must not access the network")

    monkeypatch.setattr("socket.create_connection", fail_network)
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["accounts_payable"])
    content = pack_process_map_html_bytes(pack)
    html = content.decode("utf-8")

    assert isinstance(content, bytes)
    assert "Accounts Payable &mdash; Process Map" in html
    assert escape(pack.mermaid_process_map) in html
    assert html.count(escape(pack.mermaid_process_map)) == 1
    assert 'class="diagram" id="diagram"' in html
    assert "mermaid.render('finance-process-map', getSource())" in html
    assert "mermaid.esm.min.mjs" in html
    assert all(name in html for name in ("gate", "step", "decide", "fix"))
    assert '<details class="source">' in html
    assert '<details class="source" open' not in html
    assert "Advanced: Mermaid source" in html
    assert "Copy Mermaid source" in html
    assert "Download Mermaid source" in html
    assert "function downloadSource()" in html
    assert "Diagram could not render automatically" in html
    assert "Print / Save as PDF" in html


def test_markdown_and_docx_include_visual_sections() -> None:
    pack = generate_pack(DEFAULT_SAMPLE_INPUTS["accounts_payable"])
    markdown = pack_to_markdown(pack)
    document = Document(BytesIO(pack_to_docx_bytes(pack)))
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    assert "## Visual Process Documentation" in markdown
    assert "```mermaid" in markdown
    assert "## Control-Risk Matrix" in markdown
    assert "Visual Process Documentation" in text
    assert "Control-Risk Matrix" in text


def test_generate_examples_creates_representative_matrix_samples(tmp_path: Path) -> None:
    generated = generate_examples(tmp_path)
    names = {path.name for path in generated}

    assert "accounts_payable_control_risk_matrix.csv" in names
    assert "accounts_payable_control_risk_matrix.xlsx" in names
    assert "month_end_close_control_risk_matrix.csv" in names
    assert "payroll_controls_control_risk_matrix.xlsx" in names
    assert "accounts_payable_process_map.html" in names
    assert "month_end_close_process_map.html" in names
    assert "payroll_controls_process_map.html" in names
